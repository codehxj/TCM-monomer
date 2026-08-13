from pathlib import Path
import csv

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "manuscript.docx"
LOG = ROOT / "revision_log.md"
MATRIX = ROOT / "revision_work_matrix.csv"


def set_paragraph(paragraph, text, style=None):
    paragraph.clear()
    if style:
        paragraph.style = style
    paragraph.add_run(text)


def insert_after(paragraph, text, style="Body Text"):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


doc = Document(DOCX)
p = doc.paragraphs
stage_gate_paragraphs = [p[idx] for idx in range(82, 89)]

replacements = {
    0: ("From bench evidence to clinically relevant priorities: Context-specific roles of TCM-associated isolated natural compounds in colorectal cancer", "First Paragraph"),
    7: ("Background: Evidence for natural compounds in colorectal cancer (CRC) is often obscured when purified compounds are discussed together with formulas, extracts, derivatives, delivery systems and computational predictions. This conflation weakens attribution and can make broadly reported cytotoxicity appear more clinically informative than it is. A clinically relevant synthesis therefore needs to ask which isolated compounds map to specific resistance, metastatic or immune-niche, and inflammation-associated CRC contexts while preserving preclinical evidence boundaries.", "Body Text"),
    8: ("Purpose: To map directly attributable isolated-compound evidence to defined CRC problems and identify clinically relevant priorities and context-specific roles supported by current preclinical evidence.", "Body Text"),
    9: ("Study Design: A corpus-based evidence map and problem-oriented synthesis.", "Body Text"),
    12: ("Conclusion: The evidence does not support treating TCM-associated isolated compounds as a uniform class of stand-alone cytotoxic agents. Instead, selected compounds occupy distinct clinically relevant preclinical niches: luteolin and kaempferol are most consistently linked to treatment sensitization, tanshinone IIA to chemoresistance and inflammation-associated tumorigenesis, astragaloside IV to the liver-metastatic immune niche, and honokiol to genotype-defined cetuximab sensitization. These findings remain preclinical and should guide hypothesis selection rather than claims of patient benefit.", "Body Text"),
    21: ("Existing reviews commonly organize the field by compound names or signaling pathways. The present review instead asks which directly attributable isolated-compound findings define clinically relevant priorities within specific CRC contexts. Evidence density is therefore treated as distinct from evidence maturity, and cell or animal findings are not interpreted as patient-benefit evidence.", "Body Text"),
    22: ("The article first defines material-attribution and model-evidence boundaries, then interprets retained studies through three CRC problem domains: treatment resistance and sensitization; metastatic and immune-niche progression; and inflammation-associated colorectal tumorigenesis. The aim is to clarify context-specific roles supported by current preclinical evidence, while using mechanism modules as explanatory layers rather than as proof of patient benefit.", "Body Text"),
    25: ("This article is a corpus-based evidence map and problem-oriented synthesis, not a database-wide systematic review, formal scoping review, meta-analysis or patient-benefit assessment. The quantitative audit used 303 PubMed-derived row-level records selected for relevance to CRC and isolated natural compounds. These records were checked against PubMed metadata, local source documents and supplementary extraction tables.", "First Paragraph"),
    30: ("Priority candidates were selected by problem-candidate fit rather than publication count or reported potency. Positioning required directly attributable purified-compound evidence, an animal model or a strong treatment-defined experimental context, relevance to resistance, sensitization, inflammation, metastatic progression or immune microenvironment biology, and a traceable Chinese materia medica association.", "Body Text"),
    32: ("The evidence map describes where research is concentrated and how much survives stricter attribution. Counts of compounds, pathways or records indicate research density; they do not establish causal maturity, comparative potency or clinical value.", "Body Text"),
    34: ("Of 217 retained rows, 141 carried first-pass experimental isolated-compound annotations and 76 were contextual at the row level (Figure 1). Figures 2-4 visualize chemical-class distribution, evidence grades and model use. These figures describe research distribution and evidence density, not efficacy ranking: 77 references provided direct purified-compound experimental evidence, including 29 with animal validation and 48 that were cell-only, whereas 91 references were contextual.", "Body Text"),
    39: ("Evidence density does not equal evidence maturity. Kaempferol and luteolin were frequently annotated, yet their relevance depended on resistant or immune-relevant designs rather than on record count. Conversely, astragaloside IV and honokiol had fewer records but more interpretable links to liver-metastatic immune niches and KRAS-defined cetuximab sensitization.", "Body Text"),
    41: ("Strict attribution changes the organizing question from \"which pathways are reported?\" to \"which CRC problem can the evidence make experimentally actionable?\" Figure 5 summarizes compound-mechanism themes, whereas Figure 6 shows mechanism-annotation frequency; both should be interpreted as maps of annotated themes rather than proof of causal mechanism. Table 3 introduces candidate positioning, and Table 4 distinguishes association from causal support.", "Body Text"),
    58: ("Context-Specific Evidence-Based Roles", "Heading 1"),
    59: ("Candidate-specific therapeutic relevance", "Heading 2"),
    60: ("Candidates were positioned in the setting addressed by their strongest directly attributable evidence. Table 3 records the proposed CRC context, evidence-based role, highest evidence level, principal weakness and the experiment most likely to clarify that role.", "Body Text"),
    61: ("Luteolin. Most supported CRC setting: chemotherapy or redox-related sensitization, including ferroptosis-immune coupling. Its informative evidence links GPX4-dependent ferroptosis to immune-competent MC38 experiments and includes 5-FU, oxaliplatin and erastin combinations (Q. Cao et al. 2025; Erdogan, Agca, and Askn 2022; Jang et al. 2022; Yang et al. 2024; Y. Zheng et al. 2023). The principal limitation is that exposure-matched rescue and normal-intestinal selectivity remain incomplete, so luteolin should be framed as a sensitization hypothesis rather than a stand-alone cytotoxic candidate.", "Body Text"),
    62: ("Kaempferol. Most supported CRC setting: metabolic tolerance and 5-FU or oxaliplatin sensitization. Resistant-cell and animal-linked studies implicate PKM2, PKM splicing and PFKFB4 programs (Q. Li et al. 2019; Wu et al. 2021, 2022, 2025; J. Park et al. 2021). Its evidence is stronger for treatment-context modulation than for generalized cytotoxicity, but causal metabolic rescue and exposure-confirmed combination testing remain key limitations.", "Body Text"),
    63: ("Tanshinone IIA. Most supported CRC settings: 5-FU resistance and inflammation-associated colorectal tumorigenesis. Resistant-cell/xenograft and AOM/DSS studies provide greater model depth than marker-only reports (Dong et al. 2025; Y. Cao et al. 2023; L. Liu et al. 2021). Its main value lies in these disease-defined contexts, while formulation-defined exposure, normal-colon safety and combination-toxicity data remain insufficient for patient-benefit inference.", "Body Text"),
    64: ("Astragaloside IV. Most supported CRC setting: extracellular-vesicle and macrophage-mediated modification of the liver-metastatic niche. MC38/RAW264.7 and spleen-liver experiments connect extracellular-vesicle release, M2 polarization and metastasis (Zhou et al. 2024), making this a more coherent niche-modulation hypothesis than a generic anti-metastatic claim. The key weakness is incomplete macrophage-dependency and exposure evidence.", "Body Text"),
    65: ("Honokiol. Most supported CRC setting: genotype-defined cetuximab sensitization, particularly in KRAS-mutant CRC. Animal-linked SNX3-retromer disruption supports this setting, whereas magnolol- or baicalin-containing combinations cannot establish isolated-honokiol immune sensitization (Q. Zhu et al. 2024; Gao et al. 2025). Its role should therefore be limited to a purified-compound, genotype-specific sensitization hypothesis with remaining gaps in SNX3 rescue, exposure and intestinal-toxicity assessment.", "Body Text"),
    66: ("Macrocarpal I. The supplied records support an animal-linked immunogenic-cell-death and anti-PD-1 combination hypothesis involving tubulin/PARP1 (Qi et al. 2021; Y. Zhang et al. 2025). However, the source field identifies only a generic plant or natural-compound origin and does not establish direct traceability to Chinese materia medica. Macrocarpal I is therefore retained as a related natural-product comparator, not as a core TCM-associated candidate.", "Body Text"),
    68: ("Curdione is directly traceable in the supplied records to Ezhu (Curcumae Rhizoma) and remains a TCM-associated redox-death lead. By contrast, beta-lapachone is linked to the lapacho tree and sanguinarine mainly to Sanguinaria canadensis; neither supplied source establishes a direct Chinese materia medica association. They are discussed only as related natural-product evidence. Rhein and beta-sitosterol remain context-specific mechanism leads, while formulation and nanodelivery findings for silibinin/silybin are not assigned to the free parent compound (F. Wang et al. 2023; L. Zhao et al. 2024; Pallichankandy et al. 2023; H. Zhang et al. 2021; Gu et al. 2023; Rahimnia et al. 2024).", "Body Text"),
    70: ("The central finding is not that TCM-associated isolated compounds form a single anticancer category, but that only selected compounds align with distinct CRC problem contexts. Generic cytotoxicity and pathway-marker modulation do not identify treatment bottlenecks, selectivity, exposure feasibility or a rational complement to current care.", "Body Text"),
    71: ("The most informative evidence is problem-oriented: luteolin and kaempferol align with sensitization contexts, tanshinone IIA with 5-FU resistance and inflammation-associated tumorigenesis, astragaloside IV with an extracellular-vesicle/macrophage liver-metastatic niche, and honokiol with KRAS-defined cetuximab sensitization. This pattern argues against ranking compounds by publication count, number of reported pathways or low in vitro IC50 values.", "Body Text"),
    72: ("Mechanistic associations become useful when they sharpen a clinically relevant preclinical hypothesis rather than when they accumulate additional markers. Figures 2-4 therefore describe research distribution and evidence density, whereas Figures 5-6 summarize annotated mechanism themes; neither set should be interpreted as efficacy ranking or causal validation.", "Body Text"),
    73: ("Within the curated PubMed-derived corpus, no evaluable isolated-compound clinical evidence was identified. The review therefore supports clinically informed preclinical priorities, not claims of patient benefit.", "Body Text"),
    74: ("Clinical relevance and research priorities", "Heading 1"),
    75: ("Treatment-resistance and sensitization priorities", "Heading 2"),
    76: ("For resistance and sensitization candidates, the central question is whether an interaction addresses a defined resistant state rather than adding nonspecific toxicity. The strongest priorities are luteolin and kaempferol for chemotherapy or redox-related sensitization, tanshinone IIA for 5-FU-resistant models, and honokiol for KRAS-defined cetuximab sensitization. Future work should emphasize causal rescue, treatment-naive comparators, exposure confirmation, intestinal selectivity and safety testing with 5-FU, oxaliplatin or cetuximab.", "Body Text"),
    77: ("Metastatic and immune-niche priorities", "Heading 2"),
    78: ("For metastatic and immune-niche candidates, migration assays and subcutaneous xenografts are inadequate endpoints. Astragaloside IV is the clearest priority for extracellular-vesicle/macrophage-mediated liver-metastatic niche modification, while luteolin provides a ferroptosis-immune activation hypothesis. Orthotopic or liver-metastasis models, immune-competent hosts, macrophage or CD8 dependency testing and separation of tumor-intrinsic from microenvironment-mediated effects are needed before assigning an adjunctive role.", "Body Text"),
    79: ("Inflammation-associated CRC priorities", "Heading 2"),
    80: ("For CAC-oriented candidates, dosing before tumor initiation tests prevention or tumor-promotion biology, whereas post-initiation dosing addresses established disease. Tanshinone IIA is the most coherent TCM-associated isolated-compound example in this domain, but AOM/DSS findings should not be described as treatment evidence for advanced or metastatic CRC. Future studies should align barrier integrity, histology, cytokines and normal-colon toxicity with formulation-defined dose and tissue exposure.", "Body Text"),
    81: ("Across all domains, chemical identity, purity, batch documentation and formulation definition are prerequisites. Table 5 summarizes recurrent gaps in PK, tissue exposure, normal-intestinal selectivity and systemic tolerability; carrier-enabled exposure cannot be assigned to the free parent compound without matched comparison. Rather than listing separate stage gates as the main conclusion, these gaps should guide a compact set of priorities: disease-relevant models, mechanistic rescue, normal-intestinal comparison and combination-treatment safety.", "Body Text"),
    90: ("This corpus-based evidence map is not an exhaustive systematic review and does not represent all CRC natural-product literature. The available files did not contain a reproducible PubMed query history or export date, and the quantitative corpus includes records published before the recent five-year period. No formal risk-of-bias assessment or meta-analysis was performed; material, exposure and safety reporting were heterogeneous. Within the audited corpus, evidence remained preclinical and no evaluable isolated-compound clinical study was identified. These constraints limit generalizability but make the conservative attribution framework useful for distinguishing context-specific, problem-oriented signals from indirect or mismatched evidence.", "Body Text"),
    92: ("Taken together, the most informative evidence does not support treating TCM-associated isolated compounds as a uniform class of stand-alone cytotoxic agents. Instead, selected compounds occupy distinct and clinically relevant preclinical niches: luteolin and kaempferol are most consistently linked to treatment sensitization, tanshinone IIA to chemoresistance and inflammation-associated tumorigenesis, astragaloside IV to the liver-metastatic immune niche, and honokiol to genotype-defined cetuximab sensitization. This context-specific positioning, rather than the number of reported pathways or cytotoxicity studies, is the principal conclusion of this review.", "Body Text"),
    93: ("These findings remain preclinical and should not be interpreted as evidence of patient benefit. Their value lies in narrowing future research to defined compound-disease-context pairs and in shifting the field away from repetitive pathway reporting toward treatment-relevant models, mechanistic verification and rational combination strategies. The next step is therefore not an indiscriminate expansion of bench studies, but a more clinically informed selection of hypotheses that can determine where individual TCM-associated compounds may meaningfully complement current CRC treatment research.", "Body Text"),
}

for idx, (text, style) in replacements.items():
    set_paragraph(p[idx], text, style)

# Add domain-level synthesis paragraphs in reverse order so original paragraph indices remain valid.
insert_after(p[57], "Thus, the inflammation-associated CRC domain is most useful for separating prevention and tumor-promotion biology from treatment of established or metastatic CRC. Tanshinone IIA is the clearest TCM-associated isolated-compound example in the supplied AOM/DSS-oriented evidence, whereas formononetin and rhein support narrower CAC-relevant mechanistic hypotheses. None of these data justify describing AOM/DSS effects as patient-benefit evidence for advanced CRC.", "Body Text")
insert_after(p[55], "Across the metastatic and immune-niche domain, the strongest signals are not generic anti-invasion claims but model-context links. Astragaloside IV is most directly positioned around extracellular-vesicle release, macrophage polarization and the liver-metastatic niche; luteolin illustrates ferroptosis-associated immune activation; and Macrocarpal I should remain a related natural-product comparator because its Chinese materia medica traceability is not established in the supplied data. Figures 5 and 6 should therefore be read as maps of mechanism themes and annotation density, not as proof that these mechanisms are causally validated across compounds.", "Body Text")
insert_after(p[51], "Taken together, the treatment-resistance domain supports a context-specific sensitization argument rather than a broad anti-CRC ranking. Luteolin and kaempferol are most coherently aligned with chemotherapy or redox-related sensitization, tanshinone IIA has the strongest overlap with 5-FU-resistant and CAC-linked models, and honokiol is best interpreted in a KRAS-defined cetuximab-sensitization setting. These assignments remain preclinical and should be read as compound-disease-context hypotheses, not as evidence of patient benefit.", "Body Text")

# Remove the former stand-alone stage-gate section after merging it into research priorities.
for para in stage_gate_paragraphs:
    delete_paragraph(para)

# Some body paragraphs were styled as headings in the source DOCX; normalize long prose
# paragraphs while preserving true section headings.
heading_labels = {
    "Abstract",
    "Keywords",
    "Introduction",
    "Evidence Acquisition and Curation",
    "Data sources and eligibility",
    "Evidence verification, coding and consensus",
    "Evidence Map and Research Landscape",
    "Overall evidence distribution",
    "Chemical-class distribution",
    "Evidence level and model distribution",
    "Problem-Oriented Evidence Synthesis",
    "Treatment resistance and context-specific therapeutic sensitization",
    "Cell-death regulation: distinguishing sensitization from stand-alone cytotoxicity",
    "Metabolic adaptation and therapy tolerance",
    "Tumor stemness and sustained growth as tolerance states",
    "Metastatic and immune-niche progression",
    "Immune-microenvironment remodeling and checkpoint sensitization",
    "Inflammation-associated colorectal tumorigenesis",
    "Context-Specific Evidence-Based Roles",
    "Candidate-specific therapeutic relevance",
    "Association boundaries for additional compounds",
    "Discussion",
    "Clinical relevance and research priorities",
    "Treatment-resistance and sensitization priorities",
    "Metastatic and immune-niche priorities",
    "Inflammation-associated CRC priorities",
    "Limitations",
    "Conclusion",
    "Declaration of Competing Interest",
    "Ethics Approval and Consent to Participate",
    "Funding",
    "Data Availability",
    "Declaration of generative AI and AI-assisted technologies in the writing process",
    "References",
}

for para in doc.paragraphs:
    text = " ".join(para.text.split())
    if para.style and para.style.name.startswith("Heading") and text and text not in heading_labels and len(text) > 75:
        para.style = "Body Text"

doc.save(DOCX)

matrix_rows = [
    {
        "compound": "Luteolin",
        "source_or_tcm_association": "Recorded as associated with multiple herbs/plants and TCM-linked sources in the evidence tables",
        "crc_problem_context": "Chemotherapy/redox sensitization; ferroptosis-immune coupling",
        "drug_or_treatment_context": "5-FU, oxaliplatin, erastin; immune-competent MC38 context",
        "highest_model_level": "Animal-validated isolated monomer evidence, including syngeneic MC38 in the curated set",
        "main_mechanism": "GPX4-dependent ferroptosis, lipid peroxidation, immune activation; AMPK and DNA-repair contexts",
        "intervention_or_rescue_evidence": "Ferroptosis inhibition/GPX4-linked perturbation reported in direct records; further exposure-matched rescue needed",
        "animal_validation": "Yes",
        "proposed_role": "Sensitization and ferroptosis-immune hypothesis, not stand-alone cytotoxic candidate",
        "main_evidence_limit": "Exposure, normal-intestinal selectivity and disease-context replication remain incomplete",
    },
    {
        "compound": "Kaempferol",
        "source_or_tcm_association": "Mapped to Kaempferia/rhizome and other plant sources in the evidence tables",
        "crc_problem_context": "Metabolic therapy tolerance and chemotherapy sensitization",
        "drug_or_treatment_context": "5-FU and oxaliplatin resistance/sensitization",
        "highest_model_level": "Animal-linked and cell-only resistant-model evidence",
        "main_mechanism": "PKM2/PKM splicing, miRNA-mediated metabolism and PFKFB4-associated programs",
        "intervention_or_rescue_evidence": "Mechanism-focused studies present; causal metabolic rescue remains a key limitation",
        "animal_validation": "Yes for selected non-resistance contexts; resistant-model evidence remains narrower",
        "proposed_role": "Metabolic-resistance modulator and chemotherapy-sensitization hypothesis",
        "main_evidence_limit": "Need resistant organoid/animal replication with exposure and combination-safety testing",
    },
    {
        "compound": "Tanshinone IIA",
        "source_or_tcm_association": "Salvia miltiorrhiza / Danshen",
        "crc_problem_context": "5-FU chemoresistance and inflammation-associated colorectal tumorigenesis",
        "drug_or_treatment_context": "5-FU-resistant CRC models; AOM/DSS CAC models",
        "highest_model_level": "Animal-validated isolated monomer evidence",
        "main_mechanism": "Skp2/Akt/HK2, survivin, NF-kB-related inflammation, ROS/JNK and SLC7A11-associated ferroptosis",
        "intervention_or_rescue_evidence": "Resistant-cell/xenograft and pathway-agonist contexts are recorded; formulation-defined exposure remains incomplete",
        "animal_validation": "Yes",
        "proposed_role": "Defined chemoresistance and CAC-stage intervention hypothesis",
        "main_evidence_limit": "AOM/DSS evidence should not be generalized to advanced/metastatic CRC treatment",
    },
    {
        "compound": "Astragaloside IV",
        "source_or_tcm_association": "Astragali Radix / Huangqi",
        "crc_problem_context": "Liver-metastatic immune niche",
        "drug_or_treatment_context": "Macrophage polarization and extracellular-vesicle-mediated metastatic niche; no direct chemotherapy context",
        "highest_model_level": "Animal-validated isolated monomer evidence with spleen-liver model",
        "main_mechanism": "Extracellular-vesicle release, M2-type TAM activation and liver metastasis",
        "intervention_or_rescue_evidence": "Mechanistic association reported; macrophage-dependency rescue/depletion remains incomplete",
        "animal_validation": "Yes",
        "proposed_role": "Extracellular-vesicle/macrophage-mediated niche modulator",
        "main_evidence_limit": "Requires dependency testing and exposure-measured orthotopic or liver-metastasis replication",
    },
    {
        "compound": "Honokiol",
        "source_or_tcm_association": "Magnolia officinalis / Houpo",
        "crc_problem_context": "Genotype-defined EGFR-therapy sensitization",
        "drug_or_treatment_context": "Cetuximab sensitization in KRASG13D-mutant CRC",
        "highest_model_level": "Animal-validated isolated monomer evidence",
        "main_mechanism": "SNX3-retromer disruption; separate YAP1/Ano1 contexts",
        "intervention_or_rescue_evidence": "Cetuximab-sensitization study recorded; SNX3 rescue and matched genotype panels remain priorities",
        "animal_validation": "Yes",
        "proposed_role": "KRAS-defined cetuximab-sensitization hypothesis",
        "main_evidence_limit": "Do not generalize to broad anti-CRC or anti-PD-1 activity from mixture studies",
    },
    {
        "compound": "Macrocarpal I",
        "source_or_tcm_association": "Generic plant/natural-compound source in supplied tables; direct Chinese materia medica traceability not established",
        "crc_problem_context": "Related natural-product immunogenic-cell-death comparison",
        "drug_or_treatment_context": "Anti-PD-1 combination in MC38 model",
        "highest_model_level": "Animal-validated isolated natural-product evidence",
        "main_mechanism": "Tubulin/PARP1-associated immunogenic cell death",
        "intervention_or_rescue_evidence": "Combination context reported",
        "animal_validation": "Yes",
        "proposed_role": "Comparator only, not core TCM-associated isolated compound",
        "main_evidence_limit": "Source attribution boundary prevents core candidate positioning",
    },
    {
        "compound": "beta-Lapachone",
        "source_or_tcm_association": "Lapacho/Ipe tree; not verified as Chinese materia medica in supplied data",
        "crc_problem_context": "Related redox/ferroptosis natural-product evidence",
        "drug_or_treatment_context": "Oxaliplatin-resistant cell context and separate animal ferroptosis context",
        "highest_model_level": "Animal-validated isolated natural-product evidence, but source attribution is non-TCM",
        "main_mechanism": "JNK/NCOA4 ferritinophagy; ROS/redox pathways",
        "intervention_or_rescue_evidence": "Mechanistic evidence recorded",
        "animal_validation": "Yes",
        "proposed_role": "Related comparator only",
        "main_evidence_limit": "Should not be reclassified as a core TCM-associated candidate",
    },
    {
        "compound": "Sanguinarine",
        "source_or_tcm_association": "Mainly Sanguinaria canadensis in supplied data; Chinese materia medica traceability not established",
        "crc_problem_context": "Related cell-death pathway comparison",
        "drug_or_treatment_context": "No defined standard-treatment sensitization context in supplied direct records",
        "highest_model_level": "Animal-validated isolated natural-product evidence",
        "main_mechanism": "STUB1/GPX4 ferroptosis and KEAP1-PGAM5-AIFM1 oxeiptosis reported in separate studies",
        "intervention_or_rescue_evidence": "Mechanistic evidence is pathway-specific and partly divergent",
        "animal_validation": "Yes",
        "proposed_role": "Related comparator/mechanism lead only",
        "main_evidence_limit": "Divergent cell-death labels and non-TCM source boundary",
    },
]

with MATRIX.open("w", newline="", encoding="utf-8-sig") as f:
    writer = csv.DictWriter(f, fieldnames=list(matrix_rows[0].keys()))
    writer.writeheader()
    writer.writerows(matrix_rows)

LOG.write_text(
    """# Manuscript revision log

- Title: Replaced the previous broad cytotoxicity/evidence-map title with a shorter Phytomedicine-oriented title centered on bench evidence, clinically relevant priorities and context-specific roles.
- Abstract: Revised Background, Purpose, Study Design and Conclusion to emphasize problem-oriented preclinical priorities rather than clinical translation or uniform anticancer activity; retained all audited record and evidence counts.
- Results synthesis: Added concise synthesis paragraphs at the end of the treatment-resistance, metastatic/immune-niche and inflammation-associated CRC domains. The new paragraphs identify the most relevant compound-context pairs and clarify that figures describe distribution, density or annotation frequency rather than efficacy ranking or causal proof.
- Discussion: Rewrote the opening to state the central scientific conclusion: TCM-associated isolated compounds should not be treated as a uniform anticancer class, and the strongest signals are context-specific rather than pathway-count based.
- Research priorities: Renamed and compressed the former translational-barrier/stage-gate material into a shorter Clinical relevance and research priorities section. The revised section preserves disease-relevant models, rescue, exposure, normal-intestinal comparison and combination-safety limitations while reducing repetitive stop/advance wording.
- Conclusion: Replaced the previous gate-focused conclusion with two paragraphs that list the main compound-CRC problem combinations, state the review's principal conclusion and end with a preclinical evidence boundary.
- References: No new references were added in this revision. Existing citation relationships and the reference list were preserved to avoid introducing unverified bibliography changes.
- Internal work matrix: Created revision_work_matrix.csv from the supplied evidence tables for manuscript-checking use only; it was not added as a new figure, table or supplement.
""",
    encoding="utf-8",
)


