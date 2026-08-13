from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def main() -> None:
    parser = argparse.ArgumentParser(description="Inspect DOCX paragraphs and tables.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()

    document = Document(args.docx)
    lines: list[str] = []
    lines.append(f"FILE\t{args.docx}")
    lines.append(f"PARAGRAPHS\t{len(document.paragraphs)}")
    lines.append(f"TABLES\t{len(document.tables)}")
    lines.append(f"SECTIONS\t{len(document.sections)}")
    lines.append("")
    lines.append("=== PARAGRAPHS ===")
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.replace("\t", " ").replace("\n", " | ")
        lines.append(f"P{index:04d}\t[{paragraph.style.name}]\t{text}")
    for table_index, table in enumerate(document.tables):
        lines.append("")
        lines.append(
            f"=== TABLE {table_index} rows={len(table.rows)} cols={len(table.columns)} ==="
        )
        for row_index, row in enumerate(table.rows):
            cells = [
                cell.text.replace("\t", " ").replace("\n", " | ")
                for cell in row.cells
            ]
            lines.append(f"T{table_index}R{row_index:03d}\t" + "\t".join(cells))

    output = "\n".join(lines)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(output, encoding="utf-8")
    else:
        print(output)


if __name__ == "__main__":
    main()
